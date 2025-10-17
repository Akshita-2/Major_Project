import json
import datetime
import zipfile
from io import BytesIO

# --- ReportLab for PDF Generation ---
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

# --- python-docx for Word Document Generation ---
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_enhanced_pdf_resume(resume_data, template_style="professional"):
    """
    Generates an enhanced PDF resume with multiple template options.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    # --- Template-based Styling for Hiredly ---
    if template_style == "modern":
        primary_color = colors.HexColor('#1E90FF') # Hiredly Blue
        section_font = 'Helvetica-Bold'
    elif template_style == "creative":
        primary_color = colors.HexColor('#8E44AD') # Creative Purple
        section_font = 'Helvetica-Bold'
    else:  # professional
        primary_color = colors.HexColor('#000080') # Navy Blue
        section_font = 'Helvetica-Bold'

    # --- Custom Styles ---
    name_style = ParagraphStyle('NameStyle', parent=styles['h1'], fontSize=24, textColor=primary_color, alignment=1, spaceAfter=6)
    contact_style = ParagraphStyle('ContactStyle', parent=styles['Normal'], fontSize=10, alignment=1, spaceAfter=12)
    section_style = ParagraphStyle('SectionStyle', parent=styles['h2'], fontSize=14, textColor=primary_color, fontName=section_font, spaceBefore=12, spaceAfter=6, borderBottomWidth=1, borderBottomColor=primary_color, paddingBottom=2)
    
    # FIX: Create a dedicated style for bullet points with proper indentation.
    bullet_style = ParagraphStyle('BulletStyle', parent=styles['Normal'], leftIndent=20, spaceAfter=6)
    
    # --- Build Document Story ---
    story.append(Paragraph(resume_data.get('name', 'Your Name'), name_style))
    contact_info = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')}"
    story.append(Paragraph(contact_info, contact_style))

    section_map = {
        "summary": "Professional Summary", "skills": "Core Competencies", "experience": "Professional Experience",
        "education": "Education", "projects": "Key Projects", "certifications": "Certifications"
    }

    for key, title in section_map.items():
        content = resume_data.get(key)
        if content:
            story.append(Paragraph(title.upper(), section_style))
            if isinstance(content, list):
                if key == 'skills':
                    skill_rows = [content[i:i+3] for i in range(0, len(content), 3)]
                    skills_table = Table(skill_rows, colWidths=[2.2*inch, 2.2*inch, 2.2*inch])
                    skills_table.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP')]))
                    story.append(skills_table)
                else:
                    for item in content:
                        # FIX: Use the new 'bullet_style' here.
                        story.append(Paragraph(f"• {item}", bullet_style))
            else: # For summary string
                story.append(Paragraph(content, styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer


def create_word_resume(resume_data):
    """Generates a Word document (.docx) from the resume data."""
    doc = Document()
    doc.add_heading(resume_data.get('name', 'Your Name'), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')}"
    doc.add_paragraph(contact_info).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    section_map = {
        "Professional Summary": "summary", "Core Skills": "skills", "Professional Experience": "experience",
        "Education": "education", "Key Projects": "projects", "Certifications": "certifications"
    }

    for title, key in section_map.items():
        content = resume_data.get(key)
        if content:
            doc.add_heading(title.upper(), level=1)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(content))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_html_resume(resume_data):
    """Generates a single-file HTML resume."""
    skills_html = ''.join([f'<span class="skill-tag">{skill}</span>' for skill in resume_data.get('skills', [])])
    experience_html = ''.join([f'<li>{exp}</li>' for exp in resume_data.get('experience', [])])
    
    html_template = f"""
    <!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>{resume_data.get('name', 'Resume')}</title>
    <style>body{{font-family: Arial, sans-serif; max-width: 800px; margin: auto; padding: 20px; line-height: 1.6;}} .name{{font-size: 2.5em; color: #1E90FF; text-align: center;}} .contact{{text-align: center; color: #555; margin-bottom: 20px;}} .section-title{{font-size: 1.4em; color: #1E90FF; border-bottom: 2px solid #1E90FF; padding-bottom: 5px; margin-top: 20px;}} .skill-tag{{display: inline-block; background: #1E90FF; color: white; padding: 5px 12px; border-radius: 15px; margin: 5px; font-size: 0.9em;}} ul{{padding-left: 20px;}}</style></head>
    <body>
        <div class="name">{resume_data.get('name', '')}</div>
        <div class="contact">{resume_data.get('email', '')} | {resume_data.get('phone', '')}</div>
        <div class="section-title">Professional Summary</div><p>{resume_data.get('summary', '')}</p>
        <div class="section-title">Core Skills</div><div>{skills_html}</div>
        <div class="section-title">Experience</div><ul>{experience_html}</ul>
    </body></html>
    """
    return html_template.encode('utf-8')


def create_resume_package(resume_data, template_style):
    """
    Creates a ZIP file containing the resume in multiple formats (PDF, DOCX, HTML, JSON).
    """
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # 1. Add PDF resume
        pdf_buffer = create_enhanced_pdf_resume(resume_data, template_style)
        zip_file.writestr("Hiredly_Resume.pdf", pdf_buffer.getvalue())

        # 2. Add Word resume
        word_buffer = create_word_resume(resume_data)
        zip_file.writestr("Hiredly_Resume.docx", word_buffer.getvalue())

        # 3. Add HTML resume
        html_content = create_html_resume(resume_data)
        zip_file.writestr("Hiredly_Resume.html", html_content)
        
        # 4. Add JSON data backup
        json_content = json.dumps(resume_data, indent=2)
        zip_file.writestr("resume_data.json", json_content)

        # 5. Add a README file
        readme_content = f"""
        Hiredly AI Resume Package
        =========================
        Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

        This package contains your AI-optimized resume in multiple formats.
        """
        zip_file.writestr("README.txt", readme_content.strip())
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()